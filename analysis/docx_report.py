#!/bin/env python

from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys, csv
from parameter_extraction import *

def docx_add_figure(document,filename,width,caption):
      paragraph = document.add_paragraph()
      paragraph.aligment = WD_ALIGN_PARAGRAPH.CENTER
      run = paragraph.add_run()
      run.add_picture(filename,width)
      p = paragraph._p  # this is the actual lxml element for a paragraph
      fld_xml = """
<w:p %s>
  <w:pPr>
    <w:pStyle w:val="Caption"/>
  </w:pPr>
  <w:r>
    <w:t xml:space="preserve">Figure </w:t>
  </w:r>
  <w:fldSimple w:instr=" SEQ Figure \* ARABIC "><w:r><w:rPr><w:noProof/></w:rPr><w:t>?</w:t></w:r></w:fldSimple>
  <w:r>
    <w:t xml:space="preserve">: %s</w:t>
  </w:r>
</w:p>""" % (nsdecls('w'),caption)
      fldSimple = parse_xml(fld_xml)
      p.addnext(fldSimple)

def docx_add_dbl_figure(document,filenameA,filenameB,captionA,captionB):
      paragraph = document.add_paragraph()
      paragraph.aligment = WD_ALIGN_PARAGRAPH.CENTER
      run = paragraph.add_run()
      run.add_picture(filenameA,Inches(3.0))
      run.add_picture(filenameB,Inches(3.0))      
      p = paragraph._p  # this is the actual lxml element for a paragraph
      fld_xml = """
<w:p %s>
  <w:pPr>
    <w:pStyle w:val="Caption"/>
  </w:pPr>
  <w:r>
    <w:t xml:space="preserve">Figure </w:t>
  </w:r>
  <w:fldSimple w:instr=" SEQ Figure \* ARABIC "><w:r><w:rPr><w:noProof/></w:rPr><w:t>?</w:t></w:r></w:fldSimple>
  <w:r>
    <w:t xml:space="preserve">: (left) %s (right) %s</w:t>
  </w:r>
</w:p>""" % (nsdecls('w'),captionA,captionB)
      fldSimple = parse_xml(fld_xml)
      p.addnext(fldSimple)

# Use PIL to save some image metadata
from PIL import Image
from PIL import PngImagePlugin
import os, glob

# Grab metadata
#metafile=os.path.join('./','metadata.csv')
#if not os.path.exists(metafile):
#    exit()
#f=open(metafile)
#metareader=csv.DictReader(f,delimiter=',')
#metadata=dict()
#for meta in metareader:
#    metadata[filename] = meta
      
document = Document()

#document.add_heading('Preamble', level=0)
#p = document.add_paragraph('')

#document.add_page_break()

#document.add_heading('Electrical Measurements', level=0)

# Each argument is a directory
for pth in sys.argv[1:]:
    document.add_heading(pth, level=1)
    for imgA in glob.glob('%s/*_idvg.png'%pth):
        imgB = imgA[:-9]+'_gm.png'
        imA = Image.open(imgA)
        imB = Image.open(imgB)
        docx_add_dbl_figure(document,imgA,imgB,imA.info['caption'],imB.info['caption'])
        imA.close()
        imB.close()

#document.add_page_break()

document.add_heading('Parameter Extraction', level=0)

document.add_heading('Si nfet5 (Vd=50 mV)', level=1)
###############################################################################
# Table
###############################################################################
table=document.add_table(rows=1,cols=6)
row = table.rows[0]
(W,L,VD)=(5*1e-6,15e-9,0.05)
Ids=(10e-9)/(W/float(L))
for (i,text) in enumerate(['Temperature (K)','Dose (rad)','Vth (mV) @ Ids=%g'%(Ids),'SS','peak gm (uS)']):
    row.cells[i].text = text
for (filename,T,TID) in [
        ('Bulk nFinFET E05T10 RT/Si_bulk_nFinFET_E05T10_RT_prerad.csv','293','0k'),
        ('Bulk nFinFET E05T10 RT/Si_bulk_nFinFET_E05T10_RT_10krad.csv','293','10k'),
        ('Bulk nFinFET E05T10 RT/Si_bulk_nFinFET_E05T10_RT_50krad.csv','293','50k'),
        ('Bulk nFinFET E05T10 RT/Si_bulk_nFinFET_E05T10_RT_100krad.csv','293','100k'),
        ('Bulk nFinFET E05T10 RT/Si_bulk_nFinFET_E05T10_RT_500krad.csv','293','500k'),
        ('Bulk nFinFET E05T10 RT/Si_bulk_nFinFET_E05T10_RT_1Mrad.csv','293','1M'),
        ('Bulk nFinFET E05T10 89K/Si_bulk_nFinFET_E05T10_cryo_prerad.csv','89','0k'),
        ('Bulk nFinFET E05T10 89K/Si_bulk_nFinFET_E05T10_cryo_10krad.csv','89','10k'),
        ('Bulk nFinFET E05T10 89K/Si_bulk_nFinFET_E05T10_cryo_50krad.csv','89','50k'),
        ('Bulk nFinFET E05T10 89K/Si_bulk_nFinFET_E05T10_cryo_100krad.csv','89','100k'),
        ('Bulk nFinFET E05T10 89K/Si_bulk_nFinFET_E05T10_cryo_500krad.csv','89','500k'),
        ('Bulk nFinFET E05T10 89K/Si_bulk_nFinFET_E05T10_cryo_1Mrad.csv','89','1M'),
]:
    table.add_row()
    row = table.rows[-1]
    d=read_data(filename,'VG',['ID'],where={'VD':VD})
    vstr = vth(*d,W=W,L=L)
    vstr = ('%d'%(1e3*vstr)) if vstr else '-'
    for (i,text) in enumerate([T,TID,vstr,'%d'%SS(*d),'%d'%(1e6*peak_gm(*d)[2])]):
        row.cells[i].text = text

document.add_heading('Si nfet8 (Vd=50 mV)', level=1)
###############################################################################
# Table
###############################################################################
table=document.add_table(rows=1,cols=6)
row = table.rows[0]
(W,L,VD)=(5*1e-6,40e-9,0.05)
Ids=(10e-9)/(W/float(L))
for (i,text) in enumerate(['Temperature (K)','Dose (rad)','Vth (mV) @ Ids=%g'%(Ids),'SS','peak gm (uS)']):
    row.cells[i].text = text
for (filename,T,TID) in [
        ('Bulk nFinFET E08T10 RT/Si_bulk_nFinFET_E08T10_RT_prerad.csv','293','0k'),
        ('Bulk nFinFET E08T10 RT/Si_bulk_nFinFET_E08T10_RT_10krad.csv','293','10k'),
        ('Bulk nFinFET E08T10 RT/Si_bulk_nFinFET_E08T10_RT_50krad.csv','293','50k'),
        ('Bulk nFinFET E08T10 RT/Si_bulk_nFinFET_E08T10_RT_100krad.csv','293','100k'),
        ('Bulk nFinFET E08T10 RT/Si_bulk_nFinFET_E08T10_RT_500krad.csv','293','500k'),
        ('Bulk nFinFET E08T10 RT/Si_bulk_nFinFET_E08T10_RT_1Mrad.csv','293','1M'),
        ('Bulk nFinFET E08T10 89K/Si_bulk_nFinFET_E08T10_cryo_prerad.csv','89','0k'),
        ('Bulk nFinFET E08T10 89K/Si_bulk_nFinFET_E08T10_cryo_10krad.csv','89','10k'),
        ('Bulk nFinFET E08T10 89K/Si_bulk_nFinFET_E08T10_cryo_50krad.csv','89','50k'),
        ('Bulk nFinFET E08T10 89K/Si_bulk_nFinFET_E08T10_cryo_100krad.csv','89','100k'),
        ('Bulk nFinFET E08T10 89K/Si_bulk_nFinFET_E08T10_cryo_500krad.csv','89','500k'),
        ('Bulk nFinFET E08T10 89K/Si_bulk_nFinFET_E08T10_cryo_1Mrad.csv','89','1M'),
]:
    table.add_row()
    row = table.rows[-1]
    d=read_data(filename,'VG',['ID'],where={'VD':VD})
    vstr = vth(*d,W=W,L=L)
    vstr = ('%d'%(1e3*vstr)) if vstr else '-'
    for (i,text) in enumerate([T,TID,vstr,'%d'%SS(*d),'%d'%(1e6*peak_gm(*d)[2])]):
        row.cells[i].text = text

document.add_heading('SOI nfet5 (Vd=50 mV)', level=1)
###############################################################################
# Table
###############################################################################
table=document.add_table(rows=1,cols=6)
row = table.rows[0]
(W,L,VD)=(5*1e-6,15e-9,0.05)
Ids=(10e-9)/(W/float(L))
for (i,text) in enumerate(['Temperature (K)','Dose (rad)','Vth (mV) @ Ids=%g'%(Ids),'SS','peak gm (uS)']):
    row.cells[i].text = text
for (filename,T,TID) in [
        ('SOI nFinFET E05T10 RT/Si_SOI_nFinFET_E05T10_RT_prerad.csv','293','0k'),
        ('SOI nFinFET E05T10 RT/Si_SOI_nFinFET_E05T10_RT_10krad.csv','293','10k'),
        ('SOI nFinFET E05T10 RT/Si_SOI_nFinFET_E05T10_RT_50krad.csv','293','50k'),
        ('SOI nFinFET E05T10 RT/Si_SOI_nFinFET_E05T10_RT_100krad.csv','293','100k'),
        ('SOI nFinFET E05T10 RT/Si_SOI_nFinFET_E05T10_RT_500krad.csv','293','500k'),
        ('SOI nFinFET E05T10 RT/Si_SOI_nFinFET_E05T10_RT_1Mrad.csv','293','1M'),        
        ('SOI FinFET 1.8MeV Dose NFETs E05 E08 at 90k/0krad_nfet5_90k.csv','90','0k'),
        ('SOI FinFET 1.8MeV Dose NFETs E05 E08 at 90k/50krad_nfet5_90k.csv','90','50k'),
        ('SOI FinFET 1.8MeV Dose NFETs E05 E08 at 90k/100krad_nfet5_90k.csv','90','100k'),
        ('SOI FinFET 1.8MeV Dose NFETs E05 E08 at 90k/500krad_nfet5_90k.csv','90','500k'),
#        ('SOI FinFET 1.8MeV Dose NFETs E05 E08 at 90k/1Mrad_nfet5_90k.csv','90','1M'),        
]:
    table.add_row()
    row = table.rows[-1]
    d=read_data(filename,'VG',['ID'],where={'VD':VD})
    vstr = vth(*d,W=W,L=L)
    vstr = ('%d'%(1e3*vstr)) if vstr else '-'
    for (i,text) in enumerate([T,TID,vstr,'%d'%SS(*d),'%d'%(1e6*peak_gm(*d)[2])]):
        row.cells[i].text = text

document.add_heading('SOI nfet8 (Vd=50 mV)', level=1)
###############################################################################
# Table
###############################################################################
table=document.add_table(rows=1,cols=6)
row = table.rows[0]
(W,L,VD)=(5*1e-6,40e-9,0.05)
Ids=(10e-9)/(W/float(L))
for (i,text) in enumerate(['Temperature (K)','Dose (rad)','Vth (mV) @ Ids=%g'%(Ids),'SS','peak gm (uS)']):
    row.cells[i].text = text
for (filename,T,TID) in [
        ('SOI nFinFET E08T10 RT/Si_SOI_nFinFET_E08T10_RT_prerad.csv','293','0k'),
        ('SOI nFinFET E08T10 RT/Si_SOI_nFinFET_E08T10_RT_10krad.csv','293','10k'),
        ('SOI nFinFET E08T10 RT/Si_SOI_nFinFET_E08T10_RT_50krad.csv','293','50k'),
        ('SOI nFinFET E08T10 RT/Si_SOI_nFinFET_E08T10_RT_100krad.csv','293','100k'),
        ('SOI nFinFET E08T10 RT/Si_SOI_nFinFET_E08T10_RT_500krad.csv','293','500k'),
        ('SOI nFinFET E08T10 RT/Si_SOI_nFinFET_E08T10_RT_1Mrad.csv','293','1M'),        
        ('SOI FinFET 1.8MeV Dose NFETs E05 E08 at 90k/0krad_nfet8_90k.csv','90','0k'),
        ('SOI FinFET 1.8MeV Dose NFETs E05 E08 at 90k/50krad_nfet8_90k.csv','90','50k'),
        ('SOI FinFET 1.8MeV Dose NFETs E05 E08 at 90k/100krad_nfet8_90k.csv','90','100k'),
        ('SOI FinFET 1.8MeV Dose NFETs E05 E08 at 90k/500krad_nfet8_90k.csv','90','500k'),
        ('SOI FinFET 1.8MeV Dose NFETs E05 E08 at 90k/1Mrad_nfet8_90k.csv','90','1M'),        
]:
    table.add_row()
    row = table.rows[-1]
    d=read_data(filename,'VG',['ID'],where={'VD':VD})
    vstr = vth(*d,W=W,L=L)
    vstr = ('%d'%(1e3*vstr)) if vstr else '-'
    for (i,text) in enumerate([T,TID,vstr,'%d'%SS(*d),'%d'%(1e6*peak_gm(*d)[2])]):
        row.cells[i].text = text

document.save('demo.docx')
exit()
